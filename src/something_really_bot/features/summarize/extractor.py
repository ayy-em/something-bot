"""Text extraction from PDF / DOCX / TXT / MD documents (#46).

* PDF → PyMuPDF (``fitz``). Iterates pages and concatenates ``get_text()``.
* DOCX → ``python-docx``. Concatenates paragraph text.
* TXT / MD / unknown text/* mime types → ``bytes.decode("utf-8",
  errors="replace")``.

Anything else raises :class:`UnsupportedDocumentError`.

Pillow / fitz / docx are all sync C libraries; the extractor runs them
inside ``asyncio.to_thread`` so the FastAPI event loop stays free.
"""

import asyncio
import io
from dataclasses import dataclass

from something_really_bot.logging import get_logger

_logger = get_logger(__name__)

MAX_TEXT_CHARS = 60_000  # ~15k tokens — fits comfortably under chat-model contexts.


class UnsupportedDocumentError(Exception):
    """Raised when the file's MIME type / extension isn't on the allow list."""


class DocumentExtractionError(Exception):
    """Raised when the extractor failed to parse a supported document."""


@dataclass(frozen=True)
class ExtractedDocument:
    """Output of :func:`extract`."""

    text: str
    char_count: int
    truncated: bool


_PDF_MIME_TYPES = {"application/pdf", "application/x-pdf"}
_DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",  # DOC — best-effort; .doc isn't really supported but we let docx try
}


async def extract(
    document_bytes: bytes,
    *,
    filename: str | None,
    mime_type: str | None,
) -> ExtractedDocument:
    """Extract text from ``document_bytes`` based on MIME or extension hints."""
    try:
        return await asyncio.to_thread(_extract_sync, document_bytes, filename, mime_type)
    except (UnsupportedDocumentError, DocumentExtractionError):
        raise
    except Exception as exc:  # noqa: BLE001
        # ``filename`` is a reserved LogRecord attribute; use ``source_name``
        # in the structured payload instead.
        _logger.warning(
            "summarize_extract_failed",
            extra={"exception_type": type(exc).__name__, "source_name": filename},
        )
        raise DocumentExtractionError(str(exc)) from exc


def _extract_sync(
    document_bytes: bytes,
    filename: str | None,
    mime_type: str | None,
) -> ExtractedDocument:
    kind = _classify(filename, mime_type)

    if kind == "pdf":
        text = _extract_pdf(document_bytes)
    elif kind == "docx":
        text = _extract_docx(document_bytes)
    elif kind == "text":
        text = document_bytes.decode("utf-8", errors="replace")
    else:
        raise UnsupportedDocumentError(
            f"Unsupported document type — filename={filename!r}, mime={mime_type!r}"
        )

    original_len = len(text)
    truncated = original_len > MAX_TEXT_CHARS
    if truncated:
        text = text[:MAX_TEXT_CHARS]
    return ExtractedDocument(text=text, char_count=original_len, truncated=truncated)


def _classify(filename: str | None, mime_type: str | None) -> str:
    if mime_type:
        normalized = mime_type.lower()
        if normalized in _PDF_MIME_TYPES:
            return "pdf"
        if normalized in _DOCX_MIME_TYPES:
            return "docx"
        if normalized.startswith("text/"):
            return "text"
    if filename:
        lower = filename.lower()
        if lower.endswith(".pdf"):
            return "pdf"
        if lower.endswith(".docx"):
            return "docx"
        if lower.endswith((".txt", ".md", ".markdown", ".log", ".csv")):
            return "text"
    return "unsupported"


def _extract_pdf(data: bytes) -> str:
    # Import lazily so unit tests that never touch PDFs don't pay the
    # PyMuPDF import cost.
    import fitz  # type: ignore[import-not-found]

    parts: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    return "\n".join(p.text for p in document.paragraphs if p.text)
