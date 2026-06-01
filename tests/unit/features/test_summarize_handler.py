"""Tests for the /summarize command (#46)."""

import io
from dataclasses import dataclass, field
from datetime import UTC, datetime, timedelta
from typing import Any

from pydantic import SecretStr

from something_really_bot.config import Settings
from something_really_bot.features.summarize.extractor import (
    MAX_TEXT_CHARS,
    DocumentExtractionError,
    UnsupportedDocumentError,
    extract,
)
from something_really_bot.features.summarize.handler import (
    COMMAND_NAME,
    PROMPT_TEXT,
    SummarizeHandler,
)
from something_really_bot.features.summarize.summarizer import SummarizationError
from something_really_bot.persistence import EventRecord
from something_really_bot.routing.types import BotContext
from something_really_bot.services.pending_actions import PendingAction
from something_really_bot.telegram.client import TelegramFileError, TelegramSendError
from something_really_bot.telegram.models import (
    CommandContent,
    Document,
    DocumentContent,
    PrivateMessage,
    User,
)


def _settings() -> Settings:
    return Settings.model_construct(
        telegram_webhook_secret=SecretStr("x"),
        telegram_bot_token=SecretStr("tok"),
        telegram_qa_user_ids=frozenset(),
    )


def _ctx(pending_action: PendingAction | None = None, persistence: Any = None) -> BotContext:
    return BotContext(
        settings=_settings(),
        pending_action=pending_action,
        persistence=persistence,
    )


def _command() -> PrivateMessage:
    return PrivateMessage(
        update_id=1,
        message_id=42,
        chat_id=100,
        date=1234567890,
        content=CommandContent(command=COMMAND_NAME, text=COMMAND_NAME, args=None),
        from_user=User(id=999, is_bot=False, first_name="t"),
    )


def _document(
    *,
    filename: str = "memo.txt",
    mime_type: str | None = "text/plain",
) -> PrivateMessage:
    return PrivateMessage(
        update_id=2,
        message_id=43,
        chat_id=100,
        date=1234567891,
        content=DocumentContent(
            document=Document(
                file_id="doc-id",
                file_unique_id="doc-uniq",
                file_name=filename,
                mime_type=mime_type,
                file_size=1024,
            ),
        ),
        from_user=User(id=999, is_bot=False, first_name="t"),
    )


def _pending(command: str = COMMAND_NAME) -> PendingAction:
    now = datetime.now(UTC)
    return PendingAction(
        bot_id="default",
        chat_id=100,
        user_id=999,
        command=command,
        expected_input="document",
        metadata={},
        created_at=now,
        expires_at=now + timedelta(minutes=10),
    )


@dataclass
class _FakeTelegram:
    sent: list[dict[str, Any]] = field(default_factory=list)
    download_bytes: bytes = b"hello world"
    get_file_raises: BaseException | None = None
    download_raises: BaseException | None = None
    send_raises: BaseException | None = None

    async def send_message(self, chat_id, text, *, reply_to_message_id=None, parse_mode=None):
        if self.send_raises is not None:
            raise self.send_raises
        self.sent.append(
            {
                "chat_id": chat_id,
                "text": text,
                "reply_to_message_id": reply_to_message_id,
                "parse_mode": parse_mode,
            }
        )
        return {"message_id": 1}

    async def get_file_path(self, file_id):
        if self.get_file_raises is not None:
            raise self.get_file_raises
        return "doc/x.txt"

    async def download_file(self, file_path):
        if self.download_raises is not None:
            raise self.download_raises
        return self.download_bytes


@dataclass
class _FakeGCS:
    uploads: list[dict[str, Any]] = field(default_factory=list)

    async def upload(self, *, object_key, data, content_type=None):
        self.uploads.append(
            {"object_key": object_key, "data_len": len(data), "content_type": content_type}
        )
        return f"gs://bucket/{object_key}"


@dataclass
class _FakeSummarizer:
    response: str = "TL;DR summary."
    calls: list[str] = field(default_factory=list)
    raises: BaseException | None = None

    async def summarize(self, text):
        if self.raises is not None:
            raise self.raises
        self.calls.append(text)
        return self.response


@dataclass
class _FakePending:
    set_calls: list[dict[str, Any]] = field(default_factory=list)
    clear_calls: list[dict[str, Any]] = field(default_factory=list)

    async def set(self, **kwargs):
        self.set_calls.append(kwargs)

    async def clear(self, **kwargs):
        self.clear_calls.append(kwargs)


@dataclass
class _RecordingPersistence:
    events: list[EventRecord] = field(default_factory=list)

    def record_raw_update(self, _r): ...
    def record_message(self, _r): ...
    def record_file(self, _r): ...
    def record_response(self, _r): ...

    def record_event(self, record):
        self.events.append(record)


def _build_handler(
    *,
    telegram: _FakeTelegram | None = None,
    gcs: _FakeGCS | None = None,
    summarizer: _FakeSummarizer | None = None,
    pending: _FakePending | None = None,
    scheduler=None,
):
    tg = telegram or _FakeTelegram()
    g = gcs or _FakeGCS()
    s = summarizer if summarizer is not None else _FakeSummarizer()
    p = pending or _FakePending()
    handler = SummarizeHandler(
        scheduler=scheduler or (lambda c: c.close()),
        gcs_storage_factory=lambda: g,
        telegram_client_factory=lambda: tg,
        summarizer_factory=lambda: s,
        pending_action_store_factory=lambda: p,
    )
    return handler, tg, g, s, p


# ---------------- Extractor unit tests ---------------- #


async def test_extractor_handles_text_plain() -> None:
    result = await extract(b"hello world", filename="x.txt", mime_type="text/plain")
    assert result.text == "hello world"
    assert result.truncated is False


async def test_extractor_handles_markdown_by_extension() -> None:
    result = await extract(b"# Hi", filename="readme.md", mime_type=None)
    assert "# Hi" in result.text


async def test_extractor_truncates_at_max() -> None:
    big = ("A" * (MAX_TEXT_CHARS + 100)).encode("utf-8")
    result = await extract(big, filename="big.txt", mime_type="text/plain")
    assert result.truncated is True
    assert len(result.text) == MAX_TEXT_CHARS
    assert result.char_count == MAX_TEXT_CHARS + 100


async def test_extractor_rejects_unknown_type() -> None:
    import pytest

    with pytest.raises(UnsupportedDocumentError):
        await extract(b"\x00\x01", filename="cool.zip", mime_type="application/zip")


async def test_extractor_handles_real_docx() -> None:
    docx_bytes = _build_docx("Paragraph one.\nParagraph two.")
    result = await extract(
        docx_bytes,
        filename="memo.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert "Paragraph one." in result.text
    assert "Paragraph two." in result.text


async def test_extractor_handles_real_pdf() -> None:
    pdf_bytes = _build_pdf("Hello PDF world.")
    result = await extract(pdf_bytes, filename="paper.pdf", mime_type="application/pdf")
    assert "Hello PDF world." in result.text


async def test_extractor_corrupted_pdf_raises_extraction_error() -> None:
    import pytest

    with pytest.raises(DocumentExtractionError):
        await extract(b"%PDF-not-actually-a-pdf", filename="bad.pdf", mime_type="application/pdf")


# ---------------- Handler tests ---------------- #


def test_matches_command() -> None:
    handler, *_ = _build_handler()
    assert handler.matches(_command(), _ctx()) is True


def test_matches_document_with_pending() -> None:
    handler, *_ = _build_handler()
    assert handler.matches(_document(), _ctx(_pending())) is True


def test_does_not_match_document_without_pending() -> None:
    handler, *_ = _build_handler()
    assert handler.matches(_document(), _ctx()) is False


def test_does_not_match_document_with_other_pending() -> None:
    handler, *_ = _build_handler()
    assert handler.matches(_document(), _ctx(_pending("ocr"))) is False


async def test_command_sets_pending_and_prompts() -> None:
    handler, tg, _, _, pending = _build_handler()
    await handler.handle(_command(), _ctx())
    assert pending.set_calls[0]["command"] == COMMAND_NAME
    assert pending.set_calls[0]["expected_input"] == "document"
    assert tg.sent[0]["text"] == PROMPT_TEXT


async def test_document_happy_path() -> None:
    scheduled: list[Any] = []
    persistence = _RecordingPersistence()
    handler, tg, gcs, summ, pending = _build_handler(
        telegram=_FakeTelegram(download_bytes=b"The quick brown fox."),
        summarizer=_FakeSummarizer(response="It's a sentence about a fox."),
        scheduler=lambda c: scheduled.append(c),
    )

    await handler.handle(_document(), _ctx(_pending(), persistence=persistence))
    await scheduled[0]

    # ack + final summary reply
    assert len(tg.sent) == 2
    assert "<i>It&#x27;s a sentence about a fox.</i>" in tg.sent[1]["text"]
    assert tg.sent[1]["parse_mode"] == "HTML"

    # GCS upload of the original
    assert len(gcs.uploads) == 1
    assert gcs.uploads[0]["object_key"].startswith("summarizer/100/43/doc-uniq_memo.txt")

    # Summarizer got the extracted text
    assert summ.calls == ["The quick brown fox."]

    # Pending cleared, success event
    assert len(pending.clear_calls) == 1
    assert any(e.event == "summarize_succeeded" for e in persistence.events)


async def test_oversized_document_warns_user_about_truncation() -> None:
    scheduled: list[Any] = []
    big_text = ("X" * (MAX_TEXT_CHARS + 5_000)).encode("utf-8")
    handler, tg, _, summ, _ = _build_handler(
        telegram=_FakeTelegram(download_bytes=big_text),
        summarizer=_FakeSummarizer(response="A lot of X."),
        scheduler=lambda c: scheduled.append(c),
    )

    await handler.handle(_document(), _ctx(_pending()))
    await scheduled[0]

    # Summarizer received exactly MAX_TEXT_CHARS, no more.
    assert len(summ.calls) == 1
    assert len(summ.calls[0]) == MAX_TEXT_CHARS

    # Reply includes the truncation notice.
    assert "first 60k characters" in tg.sent[1]["text"]


async def test_unsupported_document_replies_unsupported() -> None:
    scheduled: list[Any] = []
    handler, tg, _, summ, _ = _build_handler(
        telegram=_FakeTelegram(download_bytes=b"PK\x03\x04 zip body"),
        scheduler=lambda c: scheduled.append(c),
    )

    await handler.handle(
        _document(filename="archive.zip", mime_type="application/zip"),
        _ctx(_pending()),
    )
    await scheduled[0]

    assert "PDF, DOCX, TXT, and Markdown" in tg.sent[1]["text"]
    assert summ.calls == []


async def test_empty_document_replies_empty_message() -> None:
    scheduled: list[Any] = []
    handler, tg, _, summ, _ = _build_handler(
        telegram=_FakeTelegram(download_bytes=b"   \n\n  "),
        scheduler=lambda c: scheduled.append(c),
    )

    await handler.handle(_document(), _ctx(_pending()))
    await scheduled[0]

    assert "looks empty" in tg.sent[1]["text"]
    assert summ.calls == []


async def test_summarizer_failure_replies_user_error() -> None:
    scheduled: list[Any] = []
    handler, tg, _, _, _ = _build_handler(
        telegram=_FakeTelegram(download_bytes=b"some text"),
        summarizer=_FakeSummarizer(raises=SummarizationError("boom")),
        scheduler=lambda c: scheduled.append(c),
    )

    await handler.handle(_document(), _ctx(_pending()))
    await scheduled[0]

    assert "summarization service" in tg.sent[1]["text"].lower()


async def test_download_failure_replies_user_error() -> None:
    scheduled: list[Any] = []
    handler, tg, gcs, summ, _ = _build_handler(
        telegram=_FakeTelegram(get_file_raises=TelegramFileError("not ok")),
        scheduler=lambda c: scheduled.append(c),
    )

    await handler.handle(_document(), _ctx(_pending()))
    await scheduled[0]

    assert "pull that file" in tg.sent[1]["text"]
    assert not gcs.uploads
    assert summ.calls == []


async def test_missing_summarizer_replies_unavailable() -> None:
    scheduled: list[Any] = []
    tg = _FakeTelegram(download_bytes=b"text")
    handler = SummarizeHandler(
        scheduler=lambda c: scheduled.append(c),
        gcs_storage_factory=lambda: _FakeGCS(),
        telegram_client_factory=lambda: tg,
        summarizer_factory=lambda: None,
        pending_action_store_factory=lambda: _FakePending(),
    )

    await handler.handle(_document(), _ctx(_pending()))
    await scheduled[0]

    assert "isn't configured" in tg.sent[1]["text"]


async def test_send_failure_during_reply_is_swallowed() -> None:
    scheduled: list[Any] = []
    handler, _, _, _, _ = _build_handler(
        telegram=_FakeTelegram(send_raises=TelegramSendError("nope")),
        scheduler=lambda c: scheduled.append(c),
    )
    # Should not raise.
    await handler.handle(_command(), _ctx())
    if scheduled:
        await scheduled[0]


# ---------------- DOCX / PDF helpers ---------------- #


def _build_docx(text: str) -> bytes:
    from docx import Document

    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_pdf(text: str) -> bytes:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    data = doc.write()
    doc.close()
    return data
